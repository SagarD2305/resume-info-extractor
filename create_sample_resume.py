from docx import Document
from docx.shared import Inches
import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_resume():
    doc = Document()
    
    # Add header with personal information
    name = doc.add_paragraph()
    name_run = name.add_run('John Smith')
    name_run.bold = True
    name_run.font.size = Pt(16)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('Email: john.smith@email.com\n')
    contact.add_run('Phone: (555) 123-4567\n')
    contact.add_run('Location: San Francisco, CA')
    
    # Add section: Skills
    doc.add_heading('Skills', level=1)
    skills = doc.add_paragraph()
    skills.add_run('Programming Languages: Python, JavaScript, SQL\n')
    skills.add_run('Frameworks: React, Node.js, Docker\n')
    skills.add_run('Soft Skills: Project Management, Leadership, Communication')
    
    # Add section: Work Experience
    doc.add_heading('Work Experience', level=1)
    
    # Experience 1
    exp1 = doc.add_paragraph()
    exp1.add_run('Senior Software Engineer\n').bold = True
    exp1.add_run('Tech Solutions Inc. | 2020 - Present\n').italic = True
    exp1.add_run('• Led development of cloud-based applications using Python and AWS\n')
    exp1.add_run('• Managed team of 5 developers and implemented Agile methodologies\n')
    exp1.add_run('• Improved system performance by 40% through optimization')
    
    # Experience 2
    exp2 = doc.add_paragraph()
    exp2.add_run('Software Developer\n').bold = True
    exp2.add_run('Data Systems Corp | 2018 - 2020\n').italic = True
    exp2.add_run('• Developed and maintained React-based web applications\n')
    exp2.add_run('• Collaborated with cross-functional teams to deliver projects\n')
    exp2.add_run('• Implemented CI/CD pipelines using Jenkins')
    
    # Add section: Education
    doc.add_heading('Education', level=1)
    edu = doc.add_paragraph()
    edu.add_run('Bachelor of Science in Computer Science\n').bold = True
    edu.add_run('University of Technology | 2014 - 2018').italic = True
    
    # Save the document
    doc.save('sample_resume.docx')

if __name__ == '__main__':
    create_sample_resume()
